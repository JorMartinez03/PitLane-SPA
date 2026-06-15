from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── TITLE ──
title = doc.add_heading('PitLane — SPA Alpha', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── SUBTITLE ──
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Documentación Técnica de la Aplicación\nCalendario de F1, F2, F3 y Categorías Adicionales')
run.bold = True
run.font.size = Pt(13)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Versión: Alpha 0.2.0\n').bold = True
info.add_run('Fecha: 15/06/2026\n')
info.add_run('Framework: React 19 + Vite 8 + Tailwind CSS v4 + React Router\n')
info.add_run('Deploy: Vercel (https://pitlane-spa.vercel.app)')

doc.add_paragraph('═' * 60)

# ── HELPERS ──
def add_section(title_text, content_text):
    doc.add_heading(title_text, level=1)
    for line in content_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('├') or line.startswith('└') or line.startswith('│') or line == '':
            continue
        if line.startswith('- ') or line.startswith('  - '):
            doc.add_paragraph(line, style='List Bullet')
        else:
            p = doc.add_paragraph(line)

def add_sub(title_text, content_text):
    doc.add_heading(title_text, level=2)
    for line in content_text.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('- ') or line.startswith('  * '):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)

# ── 1. INTRODUCCIÓN ──
doc.add_heading('1. Introducción', level=1)
p = doc.add_paragraph()
p.add_run(
    'PitLane es una Single Page Application (SPA) en fase Alpha que muestra el '
    'calendario de carreras de Fórmula 1, F2, F3 y categorías adicionales '
    '(IndyCar, NASCAR, WEC, MotoGP) para la temporada 2026. '
    'Está diseñada con una arquitectura escalable y desacoplada, inspirada '
    'visualmente en el sitio oficial formula1.com. Utiliza React Router para '
    'la navegación entre páginas. Actualmente desplegada '
    'automáticamente en Vercel desde GitHub.'
)
doc.add_paragraph('Objetivos principales:', style='List Bullet')
for obj in [
    'Mostrar el calendario completo de carreras organizado por categoría.',
    'Destacar dinámicamente la próxima carrera según la fecha actual.',
    'Permitir filtrar entre F1, F2 y F3 sin recargar la página.',
    'Ofrecer calendarios de categorías adicionales (IndyCar, NASCAR, WEC, MotoGP).',
    'Navegación mediante React Router con rutas dedicadas.',
    'Mantener el código desacoplado para migrar fácilmente a una API real.',
    'Despliegue automático en Vercel con cada push a GitHub.',
]:
    doc.add_paragraph(obj, style='List Bullet')

# ── 2. ESTRUCTURA ──
doc.add_heading('2. Estructura del Proyecto', level=1)
doc.add_paragraph('La aplicación sigue una estructura modular con separación clara de responsabilidades:')

tree = """
SPA Alpha/
  index.html                # Punto de entrada HTML
  package.json              # Dependencias y scripts
  vite.config.js            # Configuración de Vite + Tailwind plugin
  vercel.json               # Configuración de SPA routing para Vercel
  src/
    main.jsx                # Renderiza React en el DOM
    index.css               # Estilos globales + tema Tailwind v4
    App.jsx                 # Componente raíz con React Router
    data/
      races.js              # Datos estáticos de 71 carreras (7 categorías)
    hooks/
      useCalendar.js        # Custom Hook: estado + lógica + filtrado
    components/
      CategoryCard.jsx      # Card de categoría con gradiente y enlace
      F1Dashboard.jsx       # Orquestador: layout principal y loading
      CategoryNav.jsx       # Barra de navegación F1 / F2 / F3
      Hero.jsx              # Banner de la próxima carrera
      CalendarGrid.jsx      # Grid responsivo de tarjetas
      RaceCard.jsx          # Tarjeta individual de carrera
    pages/
      HomePage.jsx          # Página principal con cards de categorías
      CalendarPage.jsx      # Página de calendario por categoría
      ExtrasPage.jsx        # Página de categorías adicionales
"""
for line in tree.strip().split('\n'):
    doc.add_paragraph(line)

# ── 3. ARQUITECTURA ──
doc.add_heading('3. Arquitectura (Separación de Conceptos)', level=1)
doc.add_paragraph('La aplicación sigue un flujo unidireccional de datos:')

doc.add_heading('3.1 Capa de Datos', level=2)
doc.add_paragraph('Archivo: src/data/races.js')
doc.add_paragraph(
    'Contiene un array de 71 objetos con la información de todas las carreras '
    'de la temporada 2026 para 7 categorías. Cada objeto tiene la siguiente estructura:'
)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
run = code.add_run(
    '{\n'
    '  id: "f1-07",                      # Identificador único\n'
    '  name: "Gran Premio de España",     # Nombre de la carrera\n'
    '  circuit: "Circuit de Barcelona",   # Nombre del circuito\n'
    '  date: "2026-06-01",               # Fecha en formato ISO\n'
    '  category: "F1",                    # Categoría: F1, F2, F3, IndyCar, NASCAR, WEC o MotoGP\n'
    '  status: "upcoming"                 # Estado calculado\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('3.2 Capa de Lógica (Custom Hook)', level=2)
doc.add_paragraph('Archivo: src/hooks/useCalendar.js')
doc.add_paragraph(
    'Este es el núcleo de la aplicación. Centraliza toda la lógica de estado '
    'y simula una petición asíncrona. Al migrar a producción, solo se modifica '
    'este archivo. Acepta un parámetro initialCategory para definir la categoría inicial.'
)
doc.add_paragraph('Funciones principales:')
for fn in [
    'Carga asíncrona simulada (setTimeout de 600ms)',
    'Cálculo dinámico del estado de cada carrera (past/next/upcoming)',
    'Filtrado reactivo por categoría usando useMemo',
    'Detección automática de la próxima carrera',
    'Limpieza de timers con useEffect return',
]:
    doc.add_paragraph(fn, style='List Bullet')

doc.add_paragraph('Valores retornados por el Hook:')
for val in [
    'races: Carreras filtradas por categoría activa',
    'allRaces: Todas las carreras sin filtrar',
    'loading: Booleano de estado de carga',
    'activeCategory: Categoría seleccionada',
    'setActiveCategory: Función para cambiar de categoría',
    'nextRace: Objeto de la próxima carrera',
    'pastRaces: Array de carreras pasadas',
    'upcomingRaces: Array de carreras futuras',
]:
    doc.add_paragraph(val, style='List Bullet')

doc.add_heading('3.3 Capa de Presentación (Componentes)', level=2)
doc.add_paragraph(
    'Componentes visuales que solo reciben props y renderizan UI. '
    'No contienen lógica de negocio ni acceso directo a datos.'
)

doc.add_heading('3.4 Enrutamiento (React Router)', level=2)
doc.add_paragraph(
    'La aplicación utiliza react-router-dom para la navegación entre páginas '
    'con las siguientes rutas:'
)
routes = [
    ('/ — HomePage', 'Página principal con cards de categorías (F1, F2, F3, Más Categorías)'),
    ('/calendar/:category — CalendarPage', 'Calendario para una categoría específica (F1, F2, F3, IndyCar, NASCAR, WEC, MotoGP)'),
    ('/extras — ExtrasPage', 'Categorías adicionales (IndyCar, NASCAR, WEC, MotoGP)'),
]
for route, desc in routes:
    p = doc.add_paragraph()
    run = p.add_run(f'{route}: ')
    run.bold = True
    p.add_run(desc)

# ── 4. COMPONENTES ──
doc.add_heading('4. Descripción de Componentes', level=1)

doc.add_heading('4.1 CategoryCard', level=2)
doc.add_paragraph('Archivo: src/components/CategoryCard.jsx')
doc.add_paragraph('Componente reutilizable de card para cada categoría:')
for item in [
    'Gradiente de fondo único por categoría (rojo F1, azul F2, verde F3, índigo IndyCar, ámbar NASCAR, púrpura WEC, naranja MotoGP)',
    'Efecto de elevación al hover (hover:-translate-y-1)',
    'Círculos decorativos con blur-3xl para profundidad',
    'Enlace de navegación al calendario correspondiente mediante React Router Link',
    'Muestra el número de carreras disponibles para esa categoría',
    'Botón "Ver calendario" con flecha animada al hover',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 HomePage (Página Principal)', level=2)
doc.add_paragraph('Archivo: src/pages/HomePage.jsx')
doc.add_paragraph('Página principal que lista todas las categorías:')
for item in [
    'Header con logo PitLane y navegación',
    'Sección "Categorías Principales" con cards de F1, F2 y F3',
    'Sección "Otras Series" con card "Más Categorías" que muestra etiquetas de IndyCar, NASCAR, WEC, MotoGP',
    'Enlace a /extras para explorar las categorías adicionales',
    'Diseño responsivo con grid adaptable',
    'Footer con créditos de fase Alpha',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 CalendarPage (Página de Calendario)', level=2)
doc.add_paragraph('Archivo: src/pages/CalendarPage.jsx')
doc.add_paragraph(
    'Extrae la categoría de los parámetros de la URL (/calendar/:category) '
    'y renderiza el componente F1Dashboard con esa categoría.'
)

doc.add_heading('4.4 ExtrasPage (Categorías Adicionales)', level=2)
doc.add_paragraph('Archivo: src/pages/ExtrasPage.jsx')
doc.add_paragraph('Página que lista las categorías adicionales disponibles:')
for item in [
    'Cards individuales para IndyCar Series, NASCAR Cup Series, WEC y MotoGP',
    'Cada card enlaza a su calendario específico (/calendar/IndyCar, etc.)',
    'Header con enlace de vuelta a la página principal',
    'Footer con créditos de fase Alpha',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.5 F1Dashboard (Orquestador)', level=2)
doc.add_paragraph('Archivo: src/components/F1Dashboard.jsx')
doc.add_paragraph(
    'Componente principal que orquesta el calendario. Consume useCalendar '
    'con la categoría recibida como prop y distribuye los datos a los componentes hijos:'
)
for item in [
    'Header sticky con logo PitLane y enlace "← Categorías" para volver atrás',
    'Barra de navegación CategoryNav para categorías principales (F1/F2/F3)',
    'Badge de categoría para categorías no principales (IndyCar, NASCAR, etc.)',
    'Loading skeleton animado mientras se cargan los datos',
    'Hero, CalendarGrid y estadísticas en el main',
    'Footer con créditos de fase Alpha',
    'Navegación entre categorías usando useNavigate de React Router',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.6 CategoryNav (Navegación de Categorías)', level=2)
doc.add_paragraph('Archivo: src/components/CategoryNav.jsx')
doc.add_paragraph('Barra de filtros interactiva con los botones F1, F2 y F3:')
for item in [
    'El botón activo se resalta con fondo rojo F1 (bg-f1-red)',
    'Los botones inactivos muestran texto plateado que se ilumina al hover',
    'Animación suave de 200ms en todas las transiciones',
    'Sombra roja (shadow-lg shadow-f1-red/30) en el botón activo',
    'Al hacer clic, navega a la ruta de la categoría seleccionada',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.7 Hero (Próxima Carrera)', level=2)
doc.add_paragraph('Archivo: src/components/Hero.jsx')
doc.add_paragraph('Banner destacado que muestra la próxima carrera del calendario:')
for item in [
    'Borde superior rojo F1 de 4px (border-t-4 border-f1-red)',
    'Fondo con degradado diagonal de carbon a black',
    'Círculos decorativos con blur-3xl y opacidad 5% para efecto de profundidad',
    'La fecha se formatea automáticamente con toLocaleDateString("es-MX")',
    'Indicador animado (animate-pulse) junto al texto de la categoría',
    'Si no hay próxima carrera (nextRace === null), no se renderiza',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.8 RaceCard (Tarjeta de Carrera)', level=2)
doc.add_paragraph('Archivo: src/components/RaceCard.jsx')
doc.add_paragraph('Tarjeta individual para cada carrera con animaciones:')
for item in [
    'Efecto Hover - Elevación (hover:-translate-y-1): la tarjeta se eleva 4px',
    'Cambio de Borde a Rojo (hover:border-f1-red)',
    'Barra Superior Roja (scale-x-0 a scale-x-100) con despliegue horizontal',
    'Transiciones suaves (duration-300 ease-out)',
    'Cambio de color del título (group-hover:text-f1-red)',
    'Indicador de estado: punto circular rojo (siguiente), plateado (próximamente), gris (finalizado)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.9 CalendarGrid (Grid de Calendario)', level=2)
doc.add_paragraph('Archivo: src/components/CalendarGrid.jsx')
for item in [
    'grid-cols-1: 1 columna en móviles',
    'sm:grid-cols-2: 2 columnas en tablets pequeñas',
    'lg:grid-cols-3: 3 columnas en desktop',
    'xl:grid-cols-4: 4 columnas en pantallas grandes',
    'gap-4: Espaciado uniforme de 16px entre tarjetas',
    'Mensaje de estado vacío si no hay carreras disponibles',
]:
    doc.add_paragraph(item, style='List Bullet')

# ── 5. ESTILOS ──
doc.add_heading('5. Estilos y Paleta de Colores', level=1)
doc.add_paragraph(
    'Los colores están definidos en src/index.css usando @theme de Tailwind CSS v4:'
)
colors = [
    ('f1-black', '#0f0f14', 'Fondo principal de la aplicación'),
    ('f1-carbon', '#15151e', 'Fondo de tarjetas y contenedores'),
    ('f1-red', '#e10600', 'Color de acento oficial de F1'),
    ('f1-gray', '#38383f', 'Bordes y separadores'),
    ('f1-silver', '#94a3b8', 'Texto secundario'),
]
for name, hex_val, desc in colors:
    doc.add_paragraph(f'{name}: {hex_val} — {desc}', style='List Bullet')

doc.add_paragraph(
    'También se definió la fuente Inter como tipografía principal importada '
    'desde Google Fonts en el index.html. '
    'Los gradientes de las CategoryCard se definen inline con Tailwind (from-{color}-500 to-{color}-700).'
)

# ── 6. FLUJO ──
doc.add_heading('6. Flujo de Datos', level=1)
doc.add_paragraph('El flujo completo de la aplicación es el siguiente:')
steps = [
    'El usuario abre la aplicación en el navegador.',
    'main.jsx monta el componente App en el DOM.',
    'App renderiza el Router con las rutas definidas.',
    'La ruta "/" muestra HomePage con las cards de categorías.',
    'El usuario hace clic en una categoría (ej. F1).',
    'React Router navega a "/calendar/F1".',
    'CalendarPage detecta el parámetro :category y renderiza F1Dashboard.',
    'F1Dashboard ejecuta useCalendar("F1").',
    'useCalendar inicia la carga asíncrona simulada.',
    'Loading skeleton se muestra durante 600ms.',
    'Los datos se cargan y se almacenan en el estado.',
    'Se filtran las carreras por la categoría activa.',
    'Se detecta la próxima carrera por fecha.',
    'Se renderizan Hero y CalendarGrid.',
    'El usuario puede cambiar de categoría (F1/F2/F3) desde el nav, o volver a inicio con "← Categorías".',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# ── 7. ESCALABILIDAD ──
doc.add_heading('7. Estrategia de Escalabilidad', level=1)
doc.add_paragraph(
    'La aplicación está diseñada para migrar fácilmente de datos locales a una API real:'
)

doc.add_heading('7.1 Migración a API real', level=2)
doc.add_paragraph(
    'Para conectar con una API real, solo es necesario modificar el archivo useCalendar.js:'
)
for item in [
    'Reemplazar el import de races.js por una llamada fetch().',
    'Cambiar el setTimeout por la promesa de fetch().',
    'Mapear la respuesta de la API al mismo formato estandarizado.',
    'Ningún componente visual necesita modificarse.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.2 Adición de nuevas categorías', level=2)
for item in [
    'Agregar nuevas carreras al array en races.js.',
    'Agregar la categoría a HomePage.jsx o ExtrasPage.jsx según corresponda.',
    'Opcional: agregar un gradiente en el objeto GRADIENTS de CategoryCard.jsx.',
    'Para que aparezca en el nav de F1Dashboard, agregarla al array MAIN_CATEGORIES.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.3 Adición de nuevas funcionalidades', level=2)
for item in [
    'Persistencia: Agregar un contexto o estado global (Zustand, Redux).',
    'Detalles de carrera: Crear una ruta con React Router.',
    'Favoritos: Agregar estado al hook o a un contexto separado.',
    'Temporadas múltiples: Modificar races.js para incluir año y filtrar.',
]:
    doc.add_paragraph(item, style='List Bullet')

# ── 8. COMANDOS ──
doc.add_heading('8. Comandos Disponibles', level=1)
commands = [
    ('npm run dev', 'Inicia el servidor de desarrollo en http://localhost:5173'),
    ('npm run build', 'Compila la aplicación para producción en /dist'),
    ('npm run preview', 'Sirve la build de producción localmente'),
    ('npm run lint', 'Ejecuta ESLint sobre el código'),
]
for cmd, desc in commands:
    p = doc.add_paragraph()
    run = p.add_run(f'{cmd}  ')
    run.bold = True
    run.font.name = 'Consolas'
    p.add_run(desc)

# ── 9. DEPENDENCIAS ──
doc.add_heading('9. Dependencias Principales', level=1)
doc.add_heading('Producción:', level=2)
for dep in [
    'react ^19.2.6             Biblioteca principal de UI',
    'react-dom ^19.2.6         Renderizado en el DOM',
    'react-router-dom ^6.x     Enrutamiento y navegación',
]:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_heading('Desarrollo:', level=2)
for dep in [
    'vite ^8.0.12            Bundler y servidor de desarrollo',
    'tailwindcss ^4.x        Framework de CSS utilitario',
    '@tailwindcss/vite ^4.x  Plugin de Tailwind para Vite',
    '@vitejs/plugin-react    Plugin de React para Vite',
    'eslint                  Linter de JavaScript',
]:
    doc.add_paragraph(dep, style='List Bullet')

# ── 10. DEPLOY ──
doc.add_heading('10. Deploy en Vercel', level=1)
doc.add_paragraph(
    'La aplicación está desplegada en Vercel con integración continua desde GitHub.'
)
doc.add_paragraph('https://pitlane-spa.vercel.app', style='List Bullet')
doc.add_paragraph()

doc.add_heading('10.1 Configuración', level=2)
for item in [
    'Se eliminó base: "/PitLane-SPA/" de vite.config.js (Vercel sirve desde la raíz).',
    'Se añadió plugin @tailwindcss/vite a vite.config.js para procesar Tailwind v4.',
    'Se creó vercel.json con rewrites para SPA routing (todas las rutas → index.html).',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Despliegue Automático', level=2)
doc.add_paragraph(
    'Vercel está conectado al repositorio de GitHub (JorMartinez03/PitLane-SPA). '
    'Cada vez que se hace push a la rama main, Vercel despliega automáticamente '
    'la nueva versión. No es necesario ejecutar ningún comando manual.'
)

doc.add_heading('10.3 Historial de Versiones', level=2)
versions = [
    ('Alpha 0.0.1 (04/06/2026)', 'Versión inicial con datos locales, 3 categorías, animaciones CSS.'),
    ('Alpha 0.1.0 (04/06/2026)', 'Configuración de Vercel, deploy automático, fix de Tailwind CSS v4.'),
    ('Alpha 0.2.0 (15/06/2026)', 'Página principal con cards de categorías, React Router, categorías adicionales (IndyCar, NASCAR, WEC, MotoGP).'),
]
for ver, desc in versions:
    p = doc.add_paragraph()
    p.add_run(f'{ver}: ').bold = True
    p.add_run(desc)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PitLane_Documentacion.docx')
doc.save(output_path)
print(f'Documento guardado: {output_path}')
